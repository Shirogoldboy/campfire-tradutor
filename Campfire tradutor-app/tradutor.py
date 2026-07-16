import sys
import os
import re
import json
import csv
import io
import subprocess
import tempfile
import hashlib
import sqlite3
from dotenv import load_dotenv
import anthropic

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')

load_dotenv()

# ─── CONFIG DO DICIONÁRIO COLABORATIVO ───────────────────────────────────────
import time
import base64
import threading
import requests

GITHUB_TOKEN    = os.getenv("GITHUB_TOKEN", "")
GITHUB_REPO     = os.getenv("GITHUB_REPO", "Shirogoldboy/campfire-dictionary")
GITHUB_API      = "https://api.github.com"
_dict_mem_cache = {}
DICT_CACHE_TTL  = 3600

TAMANHO_CHUNK = 3000
SEP = '\n[|||]\n'
MIN_LEN_BINARIO = 4
MAX_TAMANHO_ISO_ARQUIVO = 10_000_000

client = None

def inicializar_cliente(api_key=None):
    global client
    chave = api_key or os.environ.get("ANTHROPIC_API_KEY")
    if not chave:
        print("ERRO: Nenhuma chave de API fornecida.", flush=True)
        sys.exit(1)
    client = anthropic.Anthropic(api_key=chave)

# ─── CACHE DE TRADUÇÕES ───────────────────────────────────────────────────────

CACHE_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'campfire_cache.db')

def inicializar_cache():
    conn = sqlite3.connect(CACHE_PATH)
    conn.execute('''
        CREATE TABLE IF NOT EXISTS cache (
            hash TEXT PRIMARY KEY,
            original TEXT,
            traducao TEXT,
            idioma TEXT,
            criado_em TEXT DEFAULT (datetime('now'))
        )
    ''')
    conn.commit()
    conn.close()

def buscar_cache(texto, idioma, instrucao_extra=""):
    chave = hashlib.sha256(f"{texto}||{idioma}||{instrucao_extra}".encode()).hexdigest()
    try:
        conn = sqlite3.connect(CACHE_PATH)
        row = conn.execute('SELECT traducao FROM cache WHERE hash = ?', (chave,)).fetchone()
        conn.close()
        return row[0] if row else None
    except Exception:
        return None

def salvar_cache(texto, idioma, instrucao_extra, traducao):
    chave = hashlib.sha256(f"{texto}||{idioma}||{instrucao_extra}".encode()).hexdigest()
    try:
        conn = sqlite3.connect(CACHE_PATH)
        conn.execute(
            'INSERT OR REPLACE INTO cache (hash, original, traducao, idioma) VALUES (?, ?, ?, ?)',
            (chave, texto[:500], traducao, idioma)
        )
        conn.commit()
        conn.close()
    except Exception:
        pass

def stats_cache():
    try:
        conn = sqlite3.connect(CACHE_PATH)
        total = conn.execute('SELECT COUNT(*) FROM cache').fetchone()[0]
        conn.close()
        return total
    except Exception:
        return 0

inicializar_cache()

# ─── DICIONÁRIO COLABORATIVO ──────────────────────────────────────────────────

def normalizar_idioma(idioma: str) -> str:
    idioma = idioma.lower()
    mapa = {
        "português": "ptbr", "portugues": "ptbr", "portuguese": "ptbr", "pt-br": "ptbr", "ptbr": "ptbr",
        "english": "en", "inglês": "en", "ingles": "en",
        "japanese": "ja", "japonês": "ja", "japones": "ja",
        "spanish": "es", "espanhol": "es",
        "french": "fr", "francês": "fr",
        "german": "de", "alemão": "de",
        "italian": "it", "italiano": "it",
        "chinese": "zh", "chinês": "zh",
        "korean": "ko", "coreano": "ko",
        "russian": "ru", "russo": "ru",
    }
    for key, code in mapa.items():
        if key in idioma:
            return code
    return idioma[:5].replace(" ", "")

def get_dict_key(idioma_origem: str, idioma_destino: str) -> str:
    return f"{normalizar_idioma(idioma_origem)}-{normalizar_idioma(idioma_destino)}"

def detectar_idioma_texto(texto: str) -> str:
    try:
        from langdetect import detect
        return detect(texto[:2000])
    except Exception:
        return "en"

def fetch_dictionary(dict_key: str) -> dict:
    agora = time.time()
    if dict_key in _dict_mem_cache:
        dados, ts = _dict_mem_cache[dict_key]
        if agora - ts < DICT_CACHE_TTL:
            return dados
    url = f"{GITHUB_API}/repos/{GITHUB_REPO}/contents/dictionary/{dict_key}.json"
    try:
        resp = requests.get(url, headers={"Accept": "application/vnd.github.v3.raw"}, timeout=5)
        dados = resp.json() if resp.status_code == 200 else {}
    except Exception:
        dados = {}
    _dict_mem_cache[dict_key] = (dados, agora)
    return dados

def _contribuir_background(dict_key: str, novas_traducoes: dict):
    if not GITHUB_TOKEN or not novas_traducoes:
        return
    url = f"{GITHUB_API}/repos/{GITHUB_REPO}/contents/dictionary/{dict_key}.json"
    headers = {"Authorization": f"token {GITHUB_TOKEN}", "Accept": "application/vnd.github.v3+json"}
    try:
        resp = requests.get(url, headers=headers, timeout=5)
        if resp.status_code == 200:
            atual = resp.json()
            conteudo_atual = json.loads(base64.b64decode(atual["content"]).decode("utf-8"))
            sha = atual["sha"]
        else:
            conteudo_atual = {}
            sha = None
        adicionados = 0
        for k, v in novas_traducoes.items():
            if k not in conteudo_atual:
                conteudo_atual[k] = v
                adicionados += 1
        if adicionados == 0:
            return
        conteudo_bytes = json.dumps(conteudo_atual, ensure_ascii=False, indent=2).encode("utf-8")
        payload = {
            "message": f"📖 Campfire: +{adicionados} traduções ({dict_key})",
            "content": base64.b64encode(conteudo_bytes).decode("utf-8")
        }
        if sha:
            payload["sha"] = sha
        r = requests.put(url, headers=headers, json=payload, timeout=10)
        if dict_key in _dict_mem_cache:
            dados, ts = _dict_mem_cache[dict_key]
            dados.update({k: v for k, v in novas_traducoes.items() if k not in dados})
            _dict_mem_cache[dict_key] = (dados, ts)
        print(f"📖 Dicionário desktop atualizado: +{adicionados} traduções ({dict_key})", flush=True)
    except Exception as e:
        print(f"⚠️ Erro ao contribuir pro dicionário: {e}", flush=True)

def contribute_dictionary(dict_key: str, novas_traducoes: dict):
    threading.Thread(target=_contribuir_background, args=(dict_key, novas_traducoes), daemon=True).start()


# ─── TRADUÇÃO GRATUITA (sem API key) ─────────────────────────────────────────

def traduzir_com_mymemory(texto: str, idioma_origem: str, idioma_destino: str) -> str | None:
    """Tenta traduzir via MyMemory (10.000 palavras/dia grátis por IP)."""
    try:
        codigo_origem = normalizar_idioma(idioma_origem)
        codigo_destino = normalizar_idioma(idioma_destino)
        # MyMemory usa formato 'en|pt-BR'
        mapa_mymemory = {
            'ptbr': 'pt-BR', 'en': 'en', 'es': 'es', 'fr': 'fr',
            'de': 'de', 'it': 'it', 'ja': 'ja', 'ko': 'ko',
            'zh': 'zh', 'ru': 'ru', 'ar': 'ar',
        }
        orig = mapa_mymemory.get(codigo_origem, codigo_origem)
        dest = mapa_mymemory.get(codigo_destino, codigo_destino)
        url = "https://api.mymemory.translated.net/get"
        resp = requests.get(url, params={
            'q': texto[:500],  # MyMemory limita por request
            'langpair': f"{orig}|{dest}",
        }, timeout=5)
        if resp.status_code == 200:
            data = resp.json()
            if data.get('responseStatus') == 200:
                return data['responseData']['translatedText']
    except Exception:
        pass
    return None


def traduzir_com_libretranslate(texto: str, idioma_origem: str, idioma_destino: str) -> str | None:
    """Fallback pro LibreTranslate público (pode ser instável)."""
    SERVIDORES = [
        "https://translate.terraprint.co",
        "https://libretranslate.de",
    ]
    codigo_origem  = normalizar_idioma(idioma_origem)
    codigo_destino = normalizar_idioma(idioma_destino)
    # LibreTranslate usa ISO 639-1 simples
    mapa_lt = {
        'ptbr': 'pt', 'en': 'en', 'es': 'es', 'fr': 'fr',
        'de': 'de', 'it': 'it', 'ja': 'ja', 'ko': 'ko',
        'zh': 'zh', 'ru': 'ru', 'ar': 'ar',
    }
    orig = mapa_lt.get(codigo_origem, codigo_origem[:2])
    dest = mapa_lt.get(codigo_destino, codigo_destino[:2])
    for servidor in SERVIDORES:
        try:
            resp = requests.post(f"{servidor}/translate", json={
                'q': texto[:500],
                'source': orig,
                'target': dest,
                'format': 'text',
            }, timeout=5)
            if resp.status_code == 200:
                return resp.json().get('translatedText')
        except Exception:
            continue
    return None


def avaliar_qualidade_traducao(original: str, traducao: str) -> bool:
    """
    Retorna True se a tradução parece boa, False se suspeita.
    Heurísticas simples — sem gastar tokens do Claude.
    """
    if not traducao or not original:
        return False

    # 1. Ratio de tamanho muito discrepante
    ratio = len(traducao) / max(len(original), 1)
    if ratio < 0.3 or ratio > 4.0:
        return False

    # 2. Tradução igual ao original (não traduziu nada)
    if traducao.strip().lower() == original.strip().lower():
        return False

    # 3. Muitas palavras do original ainda presentes (tradução muito literal/incompleta)
    palavras_orig = set(original.lower().split())
    palavras_trad = set(traducao.lower().split())
    sobreposicao = len(palavras_orig & palavras_trad) / max(len(palavras_orig), 1)
    if sobreposicao > 0.7 and len(palavras_orig) > 3:
        return False

    # 4. Contém mensagens de erro típicas das APIs gratuitas
    erros_conhecidos = ['translation not found', 'no translation', 'error', 'quota exceeded']
    if any(e in traducao.lower() for e in erros_conhecidos):
        return False

    return True


def traduzir_segmento_inteligente(texto: str, idioma_origem: str, idioma_destino: str) -> tuple[str, str]:
    """
    Tenta traduzir um segmento na ordem:
    1. MyMemory (gratuito)
    2. LibreTranslate (gratuito, fallback)
    3. Claude (se gratuito falhou ou parece suspeito)

    Retorna (traducao, fonte) onde fonte é 'gratis', 'claude' ou 'falhou'.
    """
    # Tenta MyMemory primeiro
    trad_gratis = traduzir_com_mymemory(texto, idioma_origem, idioma_destino)

    # Fallback pro LibreTranslate se MyMemory falhou
    if not trad_gratis:
        trad_gratis = traduzir_com_libretranslate(texto, idioma_origem, idioma_destino)

    # Avalia qualidade da tradução gratuita
    if trad_gratis and avaliar_qualidade_traducao(texto, trad_gratis):
        print(f"✅ Gratuito OK: '{texto[:40]}...' → '{trad_gratis[:40]}...'", flush=True)
        return trad_gratis, 'gratis'

    # Suspeito ou falhou — usa Claude (pode lançar exceção se sem chave)
    if client is None:
        raise Exception("CLAUDE_NECESSARIO: Tradução gratuita insuficiente para este conteúdo.")

    print(f"🤖 Claude necessário: '{texto[:40]}...'", flush=True)
    trad_claude = traduzir_bloco(texto, idioma_destino)
    return trad_claude, 'claude'


def validar_com_claude_antes_dicionario(original: str, traducao_gratis: str, idioma_destino: str) -> str:
    """
    Antes de contribuir uma tradução gratuita pro dicionário,
    Claude valida se está correta. Só gasta tokens quando suspeito.
    """
    if client is None:
        return traducao_gratis  # Sem Claude, contribui como está

    prompt = (
        f"Avalie esta tradução para {idioma_destino}.\n"
        f"Original: {original}\n"
        f"Tradução: {traducao_gratis}\n\n"
        f"Se a tradução está correta e natural, responda APENAS com ela sem alterações.\n"
        f"Se estiver errada ou estranha, corrija e responda APENAS com a versão correta."
    )
    try:
        msg = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=512,
            messages=[{"role": "user", "content": prompt}]
        )
        return msg.content[0].text.strip()
    except Exception:
        return traducao_gratis



# ─── TRADUÇÃO ─────────────────────────────────────────────────────────────────

def traduzir_bloco(texto, idioma="português brasileiro coloquial", instrucao_extra=""):
    cached = buscar_cache(texto, idioma, instrucao_extra)
    if cached:
        return cached
    prompt = f"Traduza para {idioma}. Responda APENAS com o conteúdo traduzido, sem explicações. {instrucao_extra}\n\n{texto}"
    msg = client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}]
    )
    resultado = msg.content[0].text
    salvar_cache(texto, idioma, instrucao_extra, resultado)
    return resultado

def traduzir_lista(textos, instrucao_extra="", dict_key=None):
    from concurrent.futures import ThreadPoolExecutor, as_completed

    # ── Camada 1: consulta dicionário colaborativo ────────────────────────────
    dicionario = fetch_dictionary(dict_key) if dict_key else {}
    pre_traduzidos = {}
    textos_novos_idx = []

    for i, texto in enumerate(textos):
        if texto in dicionario:
            pre_traduzidos[i] = dicionario[texto]
        else:
            textos_novos_idx.append(i)

    if pre_traduzidos:
        print(f"📖 Dicionário: {len(pre_traduzidos)} segmento(s) reutilizados, "
              f"{len(textos_novos_idx)} novos.", flush=True)

    # ── Camada 2: tenta tradução gratuita (MyMemory → LibreTranslate) ─────────
    if dict_key:
        idioma_origem_layer2 = dict_key.split('-')[0]
    elif textos_novos_idx:
        idioma_origem_layer2 = detectar_idioma_texto(' '.join(textos[i] for i in textos_novos_idx[:10]))
    else:
        idioma_origem_layer2 = None

    resultado_gratis = {}  # posição em textos_novos_idx -> tradução

    if textos_novos_idx and idioma_origem_layer2:
        def tentar_gratis(pos, texto):
            trad = traduzir_com_mymemory(texto, idioma_origem_layer2, _idioma_global)
            if not trad:
                trad = traduzir_com_libretranslate(texto, idioma_origem_layer2, _idioma_global)
            if trad and avaliar_qualidade_traducao(texto, trad):
                return pos, trad
            return pos, None

        with ThreadPoolExecutor(max_workers=8) as executor:
            futuros = [executor.submit(tentar_gratis, pos, textos[i])
                       for pos, i in enumerate(textos_novos_idx)]
            for futuro in as_completed(futuros):
                pos, trad = futuro.result()
                if trad:
                    resultado_gratis[pos] = trad

        if resultado_gratis:
            print(f"✅ Gratuito: {len(resultado_gratis)}/{len(textos_novos_idx)} segmento(s) "
                  f"traduzido(s) sem tokens.", flush=True)

    # ── Camada 3: o que sobrou (sem tradução gratuita) vai pro Claude, em batches ──
    indices_para_claude = [pos for pos in range(len(textos_novos_idx)) if pos not in resultado_gratis]
    textos_para_traduzir = [textos[textos_novos_idx[pos]] for pos in indices_para_claude]

    batches, batch_idx, batch_txt, chars = [], [], [], 0
    for i, texto in enumerate(textos_para_traduzir):
        if chars + len(texto) > TAMANHO_CHUNK and batch_idx:
            batches.append((batch_idx, batch_txt))
            batch_idx, batch_txt, chars = [i], [texto], len(texto)
        else:
            batch_idx.append(i)
            batch_txt.append(texto)
            chars += len(texto)
    if batch_idx:
        batches.append((batch_idx, batch_txt))

    traduzidos_claude = [''] * len(textos_para_traduzir)
    total = len(batches)
    concluidos = [0]

    def processar_batch(b_num, indices, lote):
        resultado = traduzir_bloco(
            SEP.join(lote),
            _idioma_global,
            f"PRESERVE os separadores [|||] exatamente onde estão. {instrucao_extra}"
        )
        partes = resultado.split('[|||]')
        concluidos[0] += 1
        print(f"Batch {concluidos[0]}/{total}...", flush=True)
        return indices, partes, lote

    MAX_PARALELO = 5
    if batches:
        with ThreadPoolExecutor(max_workers=MAX_PARALELO) as executor:
            futuros = {
                executor.submit(processar_batch, b_num, indices, lote): b_num
                for b_num, (indices, lote) in enumerate(batches, 1)
            }
            for futuro in as_completed(futuros):
                try:
                    indices, partes, lote = futuro.result()
                    for i, idx in enumerate(indices):
                        traduzidos_claude[idx] = partes[i].strip() if i < len(partes) else textos_para_traduzir[idx]
                except Exception as e:
                    print(f"AVISO: erro num batch: {e}", flush=True)

    # ── Reconstrói na ordem de textos_novos_idx ────────────────────────────────
    traduzidos_novos = [''] * len(textos_novos_idx)
    for pos, trad in resultado_gratis.items():
        traduzidos_novos[pos] = trad
    for local_i, pos in enumerate(indices_para_claude):
        traduzidos_novos[pos] = traduzidos_claude[local_i]

    # ── Mescla dicionário + novos ──────────────────────────────────────────────
    traduzidos = [''] * len(textos)
    for i, trad in pre_traduzidos.items():
        traduzidos[i] = trad
    for pos, i in enumerate(textos_novos_idx):
        traduzidos[i] = traduzidos_novos[pos]

    # ── Contribui pro dicionário em background — SÓ traduções do Claude ───────
    # MyMemory proíbe republicar sua "Public Data" (segmentos crus) em outro
    # repositório (ver Termos do MyMemory). Por isso traduções gratuitas
    # (MyMemory/LibreTranslate) são usadas no resultado do usuário mas NUNCA
    # contribuídas ao dicionário público — só as geradas pelo próprio Claude.
    if dict_key and indices_para_claude:
        def contribuir_claude():
            novas = {}
            for local_i, pos in enumerate(indices_para_claude):
                i = textos_novos_idx[pos]
                original, trad = textos[i], traduzidos_claude[local_i]
                if trad:
                    novas[original] = trad
            if novas:
                contribute_dictionary(dict_key, novas)
        threading.Thread(target=contribuir_claude, daemon=True).start()

    return traduzidos

def processar_txt(caminho):
    with open(caminho, 'r', encoding='utf-8') as f:
        texto = f.read()

    idioma_origem = detectar_idioma_texto(texto)
    dict_key = get_dict_key(idioma_origem, _idioma_global)
    print(f"📖 Idioma detectado: {idioma_origem} → dicionário: {dict_key}", flush=True)

    # Divide em parágrafos pra aproveitar o dicionário por segmento
    paragrafos = [p.strip() for p in texto.split('\n') if p.strip()]

    print(f"Traduzindo {len(paragrafos)} parágrafo(s)...", flush=True)
    traduzidos = traduzir_lista(paragrafos, dict_key=dict_key)

    return '\n\n'.join(traduzidos), '.txt'


def processar_srt(caminho):
    with open(caminho, 'r', encoding='utf-8') as f:
        conteudo = f.read()
    estrutura = []
    for bloco in re.split(r'\n\n+', conteudo.strip()):
        linhas = bloco.strip().split('\n')
        if len(linhas) >= 2 and re.match(r'^\d+$', linhas[0].strip()):
            estrutura.append({'numero': linhas[0].strip(), 'tempo': linhas[1].strip(), 'texto': '\n'.join(linhas[2:])})

    # Detecta idioma de origem pra usar o dicionário certo
    amostra = ' '.join(e['texto'] for e in estrutura[:10])
    idioma_origem = detectar_idioma_texto(amostra)
    dict_key = get_dict_key(idioma_origem, _idioma_global)
    print(f"📖 Idioma detectado: {idioma_origem} → dicionário: {dict_key}", flush=True)

    traduzidos = traduzir_lista([e['texto'] for e in estrutura], dict_key=dict_key)
    for entry, trad in zip(estrutura, traduzidos):
        entry['texto'] = trad
    return '\n\n'.join([f"{e['numero']}\n{e['tempo']}\n{e['texto']}" for e in estrutura]), '.srt'

def processar_json(caminho):
    with open(caminho, 'r', encoding='utf-8') as f:
        dados = json.load(f)
    paths, textos = [], []
    def coletar(obj, path):
        if isinstance(obj, dict):
            for k, v in obj.items(): coletar(v, path + [k])
        elif isinstance(obj, list):
            for i, v in enumerate(obj): coletar(v, path + [i])
        elif isinstance(obj, str) and obj.strip():
            paths.append(path); textos.append(obj)
    def definir(obj, path, valor):
        for chave in path[:-1]: obj = obj[chave]
        obj[path[-1]] = valor
    coletar(dados, [])
    idioma_origem = detectar_idioma_texto(' '.join(textos[:20]))
    dict_key = get_dict_key(idioma_origem, _idioma_global)
    print(f"📖 Idioma detectado: {idioma_origem} → dicionário: {dict_key}", flush=True)
    traduzidos = traduzir_lista(textos, dict_key=dict_key)
    for path, trad in zip(paths, traduzidos):
        definir(dados, path, trad)
    return json.dumps(dados, ensure_ascii=False, indent=2), '.json'

def processar_xml(caminho):
    with open(caminho, 'r', encoding='utf-8') as f:
        conteudo = f.read()
    matches = [m for m in re.finditer(r'>([^<>\n]+)<', conteudo) if m.group(1).strip()]
    textos = [m.group(1) for m in matches]
    idioma_origem = detectar_idioma_texto(' '.join(textos[:20]))
    dict_key = get_dict_key(idioma_origem, _idioma_global)
    print(f"📖 Idioma detectado: {idioma_origem} → dicionário: {dict_key}", flush=True)
    traduzidos = traduzir_lista(textos, dict_key=dict_key)
    resultado, offset = list(conteudo), 0
    for m, trad in zip(matches, traduzidos):
        s, e = m.start(1) + offset, m.end(1) + offset
        resultado[s:e] = list(trad)
        offset += len(trad) - len(m.group(1))
    return ''.join(resultado), '.xml'

def processar_csv(caminho):
    with open(caminho, 'r', encoding='utf-8') as f:
        linhas = list(csv.reader(f))
    celulas = [(i, j, c) for i, l in enumerate(linhas) for j, c in enumerate(l)
               if c.strip() and not c.replace('.','').replace(',','').replace('-','').isnumeric()]
    textos = [c for _, _, c in celulas]
    idioma_origem = detectar_idioma_texto(' '.join(textos[:20]))
    dict_key = get_dict_key(idioma_origem, _idioma_global)
    print(f"📖 Idioma detectado: {idioma_origem} → dicionário: {dict_key}", flush=True)
    traduzidos = traduzir_lista(textos, dict_key=dict_key)
    for (i, j, _), trad in zip(celulas, traduzidos):
        linhas[i][j] = trad
    output = io.StringIO()
    csv.writer(output).writerows(linhas)
    return output.getvalue(), '.csv'

def processar_pdf(caminho):
    import pdfplumber
    from fpdf import FPDF

    with pdfplumber.open(caminho) as pdf:
        paginas = [p.extract_text() or "" for p in pdf.pages]

    # ── Fallback Tesseract pra páginas sem texto (PDFs escaneados) ────────────
    paginas_sem_texto = [i for i, t in enumerate(paginas) if not t.strip()]
    if paginas_sem_texto:
        print(f"📷 {len(paginas_sem_texto)} página(s) escaneada(s). Usando Tesseract OCR...", flush=True)
        try:
            import pytesseract
            from PIL import Image as PILImage
            import fitz  # PyMuPDF

            pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
            doc_fitz = fitz.open(caminho)

            for i in paginas_sem_texto:
                page = doc_fitz[i]
                pix = page.get_pixmap(dpi=200)
                img = PILImage.frombytes("RGB", [pix.width, pix.height], pix.samples)
                texto_ocr = pytesseract.image_to_string(img, lang='por+eng+jpn')
                if texto_ocr.strip():
                    paginas[i] = texto_ocr
                    print(f"  OCR página {i+1}: {len(texto_ocr)} caracteres extraídos.", flush=True)
            doc_fitz.close()
        except Exception as e:
            print(f"⚠️ Tesseract falhou: {e}", flush=True)

    indices = [i for i, t in enumerate(paginas) if t.strip()]
    amostra = ' '.join(paginas[i] for i in indices[:3])
    idioma_origem = detectar_idioma_texto(amostra)
    dict_key = get_dict_key(idioma_origem, _idioma_global)
    print(f"📖 Idioma detectado: {idioma_origem} → dicionário: {dict_key}", flush=True)

    traduzidos = traduzir_lista([paginas[i] for i in indices], dict_key=dict_key)
    for idx, trad in zip(indices, traduzidos):
        paginas[idx] = trad

    pdf_out = FPDF()
    pdf_out.add_font("Base", "", r"C:\Windows\Fonts\arial.ttf")
    pdf_out.set_font("Base", size=11)

    for texto in paginas:
        pdf_out.add_page()
        if texto.strip():
            pdf_out.multi_cell(0, 6, texto)

    return bytes(pdf_out.output()), '.pdf'

def processar_mp3(caminho):
    from faster_whisper import WhisperModel

    print("Carregando modelo de transcrição...", flush=True)
    modelo = WhisperModel("small", device="cpu", compute_type="int8")

    print("Transcrevendo áudio...", flush=True)
    segmentos, info = modelo.transcribe(caminho)
    segmentos = list(segmentos)

    def formatar_tempo(segundos):
        h = int(segundos // 3600)
        m = int((segundos % 3600) // 60)
        s = segundos % 60
        return f"{h:02d}:{m:02d}:{s:06.3f}".replace('.', ',')

    textos = [seg.text.strip() for seg in segmentos]
    print(f"Transcrição concluída ({info.language}): {len(textos)} trechos. Traduzindo...", flush=True)

    # Whisper já detecta o idioma — usa direto pro dicionário
    dict_key = get_dict_key(info.language, _idioma_global)
    print(f"📖 Idioma do áudio: {info.language} → dicionário: {dict_key}", flush=True)

    instrucao_audio = (
        f"Este é um áudio transcrito, pode conter gírias, expressões informais e coloquialismos. "
        f"Traduza para {_idioma_global}, preservando o tom e a naturalidade do áudio original. "
        f"Se houver expressões idiomáticas, use o equivalente natural em {_idioma_global}, não traduza literalmente. "
        f"Responda apenas com o texto traduzido, sem comentários."
    )

    traduzidos = traduzir_lista(textos, instrucao_audio, dict_key=dict_key)

    linhas = []
    for i, (seg, trad) in enumerate(zip(segmentos, traduzidos), 1):
        inicio = formatar_tempo(seg.start)
        fim    = formatar_tempo(seg.end)
        linhas.append(f"{i}\n{inicio} --> {fim}\n{trad}")

    return '\n\n'.join(linhas), '.srt'

def processar_video(caminho):
    pasta = os.path.dirname(caminho)
    nome_base = os.path.splitext(os.path.basename(caminho))[0]
    srt_temp = os.path.join(pasta, f"{nome_base}_temp.srt")
    print("Extraindo legendas do vídeo...", flush=True)
    subprocess.run(['ffmpeg', '-y', '-i', caminho, '-map', '0:s:0', srt_temp],
                   capture_output=True, text=True)
    if not os.path.exists(srt_temp) or os.path.getsize(srt_temp) == 0:
        if os.path.exists(srt_temp):
            os.remove(srt_temp)
        raise Exception("Nenhuma legenda embutida encontrada no arquivo de vídeo.")
    print("Legenda extraída! Iniciando tradução...", flush=True)
    conteudo, ext = processar_srt(srt_temp)
    os.remove(srt_temp)
    return conteudo, '.srt'

def processar_epub(caminho):
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup, Comment

    book = epub.read_epub(caminho)
    TAGS_IGNORADAS = {'script', 'style'}

    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        soup = BeautifulSoup(item.get_content(), 'html.parser')
        nodes = [
            n for n in soup.find_all(string=True)
            if not isinstance(n, Comment) and n.parent.name not in TAGS_IGNORADAS and n.strip()
        ]
        if not nodes:
            continue

        textos = [str(n) for n in nodes]
        idioma_origem = detectar_idioma_texto(' '.join(textos[:20]))
        dict_key = get_dict_key(idioma_origem, _idioma_global)
        print(f"📖 [{item.get_name()}] Idioma: {idioma_origem} → dicionário: {dict_key}", flush=True)
        traduzidos = traduzir_lista(textos, dict_key=dict_key)

        for node, trad in zip(nodes, traduzidos):
            node.replace_with(trad)
        item.set_content(str(soup).encode('utf-8'))

    pasta = os.path.dirname(caminho)
    nome_base = os.path.splitext(os.path.basename(caminho))[0]
    temp_path = os.path.join(pasta, f"{nome_base}_epub_temp.epub")
    epub.write_epub(temp_path, book)

    with open(temp_path, 'rb') as f:
        conteudo = f.read()
    os.remove(temp_path)
    return conteudo, '.epub'

def processar_docx(caminho):
    from docx import Document

    doc = Document(caminho)

    def coletar_paragrafos(paragrafos):
        return [p for p in paragrafos if p.text.strip()]

    paragrafos_alvo = list(coletar_paragrafos(doc.paragraphs))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragrafos_alvo += coletar_paragrafos(cell.paragraphs)
    for section in doc.sections:
        for parte in (section.header, section.footer):
            paragrafos_alvo += coletar_paragrafos(parte.paragraphs)

    textos = [p.text for p in paragrafos_alvo]
    idioma_origem = detectar_idioma_texto(' '.join(textos[:20]))
    dict_key = get_dict_key(idioma_origem, _idioma_global)
    print(f"📖 Idioma detectado: {idioma_origem} → dicionário: {dict_key}", flush=True)

    print(f"Traduzindo {len(textos)} parágrafo(s)...", flush=True)
    traduzidos = traduzir_lista(textos, dict_key=dict_key)

    for p, trad in zip(paragrafos_alvo, traduzidos):
        if p.runs:
            p.runs[0].text = trad
            for run in p.runs[1:]:
                run.text = ''
        else:
            p.text = trad

    pasta = os.path.dirname(caminho)
    nome_base = os.path.splitext(os.path.basename(caminho))[0]
    temp_path = os.path.join(pasta, f"{nome_base}_docx_temp.docx")
    doc.save(temp_path)

    with open(temp_path, 'rb') as f:
        conteudo = f.read()
    os.remove(temp_path)
    return conteudo, '.docx'

def processar_xlsx(caminho):
    import openpyxl

    wb = openpyxl.load_workbook(caminho)
    celulas = []
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                valor = cell.value
                if isinstance(valor, str) and valor.strip() and not valor.startswith('='):
                    celulas.append(cell)

    textos = [c.value for c in celulas]
    idioma_origem = detectar_idioma_texto(' '.join(textos[:20]))
    dict_key = get_dict_key(idioma_origem, _idioma_global)
    print(f"📖 Idioma detectado: {idioma_origem} → dicionário: {dict_key}", flush=True)

    print(f"Traduzindo {len(textos)} célula(s)...", flush=True)
    traduzidos = traduzir_lista(textos, dict_key=dict_key)

    for cell, trad in zip(celulas, traduzidos):
        cell.value = trad

    pasta = os.path.dirname(caminho)
    nome_base = os.path.splitext(os.path.basename(caminho))[0]
    temp_path = os.path.join(pasta, f"{nome_base}_xlsx_temp.xlsx")
    wb.save(temp_path)

    with open(temp_path, 'rb') as f:
        conteudo = f.read()
    os.remove(temp_path)
    return conteudo, '.xlsx'

def processar_po(caminho):
    import polib

    po = polib.pofile(caminho)
    entradas = [e for e in po if e.msgid.strip() and not e.obsolete]
    singulares = [e for e in entradas if not e.msgid_plural]
    plurais    = [e for e in entradas if e.msgid_plural]

    amostra = ' '.join(e.msgid for e in (singulares + plurais)[:20])
    idioma_origem = detectar_idioma_texto(amostra) if amostra else 'en'
    dict_key = get_dict_key(idioma_origem, _idioma_global)
    print(f"📖 Idioma detectado: {idioma_origem} → dicionário: {dict_key}", flush=True)
    print(f"Traduzindo {len(entradas)} entrada(s)...", flush=True)

    if singulares:
        traduzidos = traduzir_lista([e.msgid for e in singulares], dict_key=dict_key)
        for e, trad in zip(singulares, traduzidos):
            e.msgstr = trad

    if plurais:
        trad_sing = traduzir_lista([e.msgid for e in plurais], dict_key=dict_key)
        trad_plur = traduzir_lista([e.msgid_plural for e in plurais], dict_key=dict_key)
        for e, ts, tp in zip(plurais, trad_sing, trad_plur):
            e.msgstr_plural = {'0': ts, '1': tp}

    pasta = os.path.dirname(caminho)
    nome_base = os.path.splitext(os.path.basename(caminho))[0]
    temp_path = os.path.join(pasta, f"{nome_base}_po_temp.po")
    po.save(temp_path)

    with open(temp_path, 'r', encoding='utf-8') as f:
        conteudo = f.read()
    os.remove(temp_path)
    return conteudo, '.po'


def _unescape_strings_valor(s):
    return s.replace('\\"', '"').replace('\\n', '\n').replace('\\\\', '\\')

def _escape_strings_valor(s):
    return s.replace('\\', '\\\\').replace('\n', '\\n').replace('"', '\\"')

def processar_strings(caminho):
    with open(caminho, 'r', encoding='utf-8') as f:
        conteudo = f.read()

    pattern = re.compile(r'"((?:\\.|[^"\\])*)"\s*=\s*"((?:\\.|[^"\\])*)"\s*;')
    matches = list(pattern.finditer(conteudo))
    if not matches:
        raise Exception("Nenhum par chave/valor encontrado neste arquivo .strings.")

    valores = [_unescape_strings_valor(m.group(2)) for m in matches]
    idioma_origem = detectar_idioma_texto(' '.join(valores[:20]))
    dict_key = get_dict_key(idioma_origem, _idioma_global)
    print(f"📖 Idioma detectado: {idioma_origem} → dicionário: {dict_key}", flush=True)
    print(f"Traduzindo {len(matches)} entrada(s)...", flush=True)

    traduzidos = traduzir_lista(valores, dict_key=dict_key)

    resultado, offset = list(conteudo), 0
    for m, trad in zip(matches, traduzidos):
        s, e = m.start(2) + offset, m.end(2) + offset
        trad_escapado = _escape_strings_valor(trad)
        resultado[s:e] = list(trad_escapado)
        offset += len(trad_escapado) - len(m.group(2))
    return ''.join(resultado), '.strings'


def processar_resx(caminho):
    import xml.etree.ElementTree as ET

    tree = ET.parse(caminho)
    root = tree.getroot()

    entradas = []
    for data in root.findall('data'):
        if data.get('type'):
            continue
        value_el = data.find('value')
        if value_el is not None and value_el.text and value_el.text.strip():
            entradas.append(value_el)

    if not entradas:
        raise Exception("Nenhuma entrada de texto traduzível encontrada neste arquivo .resx.")

    textos = [el.text for el in entradas]
    idioma_origem = detectar_idioma_texto(' '.join(textos[:20]))
    dict_key = get_dict_key(idioma_origem, _idioma_global)
    print(f"📖 Idioma detectado: {idioma_origem} → dicionário: {dict_key}", flush=True)
    print(f"Traduzindo {len(entradas)} entrada(s)...", flush=True)

    traduzidos = traduzir_lista(textos, dict_key=dict_key)
    for el, trad in zip(entradas, traduzidos):
        el.text = trad

    pasta = os.path.dirname(caminho)
    nome_base = os.path.splitext(os.path.basename(caminho))[0]
    temp_path = os.path.join(pasta, f"{nome_base}_resx_temp.resx")
    tree.write(temp_path, encoding='utf-8', xml_declaration=True)

    with open(temp_path, 'r', encoding='utf-8') as f:
        conteudo = f.read()
    os.remove(temp_path)
    return conteudo, '.resx'

def detectar_texto_8bit(data):
    pattern = re.compile(rb'[\x20-\x7E]{%d,}' % MIN_LEN_BINARIO)
    candidatos = []
    for m in pattern.finditer(bytes(data)):
        texto = m.group().decode('ascii')
        letras = sum(c.isalpha() for c in texto)
        if letras < max(3, len(texto) * 0.5):
            continue
        candidatos.append((m.start(), m.end(), texto, '8bit'))
    return candidatos

def detectar_texto_16bit(data):
    pattern = re.compile(rb'(?:[\x20-\x7E]\x00){%d,}' % MIN_LEN_BINARIO)
    candidatos = []
    for m in pattern.finditer(bytes(data)):
        bruto = m.group()
        texto = bruto[0::2].decode('ascii')
        letras = sum(c.isalpha() for c in texto)
        if letras < max(3, len(texto) * 0.5):
            continue
        candidatos.append((m.start(), m.end(), texto, '16bit'))
    return candidatos

def detectar_blocos_comprimidos(data):
    import zlib
    achados = []
    cabecalhos = [b'\x78\x01', b'\x78\x5e', b'\x78\x9c', b'\x78\xda']
    bruto = bytes(data)
    pos = 0
    while True:
        proximo = None
        for cab in cabecalhos:
            idx = bruto.find(cab, pos)
            if idx != -1 and (proximo is None or idx < proximo):
                proximo = idx
        if proximo is None:
            break
        try:
            descomprimido = zlib.decompress(bruto[proximo:proximo + 200000])
            trechos = re.findall(rb'[\x20-\x7E]{6,}', descomprimido)
            if trechos:
                preview = trechos[0][:60].decode('ascii', errors='ignore')
                achados.append((proximo, preview))
        except Exception:
            pass
        pos = proximo + 2
    return achados

# ─── DESCOMPRESSÃO NINTENDO (LZ10/LZ11) ──────────────────────────────────────

def descomprimir_lz10(data: bytes) -> bytes | None:
    """LZ10: compressão Nintendo DS/GBA. Magic byte = 0x10."""
    try:
        if len(data) < 4 or data[0] != 0x10: return None
        tam = int.from_bytes(data[1:4], 'little')
        out, pos = bytearray(), 4
        while len(out) < tam and pos < len(data):
            flags = data[pos]; pos += 1
            for bit in range(7, -1, -1):
                if len(out) >= tam or pos >= len(data): break
                if flags & (1 << bit):
                    if pos + 1 >= len(data): break
                    b1, b2 = data[pos], data[pos+1]; pos += 2
                    dist, length = ((b1 & 0xF) << 8 | b2) + 1, (b1 >> 4) + 3
                    for _ in range(length): out.append(out[-dist] if len(out) >= dist else 0)
                else:
                    out.append(data[pos]); pos += 1
        return bytes(out)
    except Exception: return None

def descomprimir_lz11(data: bytes) -> bytes | None:
    """LZ11: variante estendida do LZ10. Magic byte = 0x11."""
    try:
        if len(data) < 4 or data[0] != 0x11: return None
        tam = int.from_bytes(data[1:4], 'little')
        out, pos = bytearray(), 4
        while len(out) < tam and pos < len(data):
            flags = data[pos]; pos += 1
            for bit in range(7, -1, -1):
                if len(out) >= tam or pos >= len(data): break
                if flags & (1 << bit):
                    if pos >= len(data): break
                    ind = data[pos]; pos += 1
                    if ind >> 4 == 1:
                        if pos + 1 >= len(data): break
                        b2, b3 = data[pos], data[pos+1]; pos += 2
                        length = ((ind & 0xF) << 12 | b2 << 4 | b3 >> 4) + 0x111
                        dist = ((b3 & 0xF) << 8 | data[pos]) + 1; pos += 1
                    elif ind >> 4 == 0:
                        if pos >= len(data): break
                        b2 = data[pos]; pos += 1
                        length = ((ind & 0xF) << 4 | b2 >> 4) + 0x11
                        dist = ((b2 & 0xF) << 8 | data[pos]) + 1; pos += 1
                    else:
                        if pos >= len(data): break
                        b2 = data[pos]; pos += 1
                        length, dist = (ind >> 4) + 1, ((ind & 0xF) << 8 | b2) + 1
                    for _ in range(length): out.append(out[-dist] if len(out) >= dist else 0)
                else:
                    out.append(data[pos]); pos += 1
        return bytes(out)
    except Exception: return None

def tentar_descomprimir_nintendo(data: bytes) -> bytes | None:
    """Tenta LZ10, LZ11 e zlib em sequência."""
    if not data: return None
    r = descomprimir_lz10(data)
    if r: return r
    r = descomprimir_lz11(data)
    if r: return r
    import zlib
    for wb in (-15, 15, 47):
        try: return zlib.decompress(data, wb)
        except Exception: pass
    return None

# ─── DETECÇÃO EUC-JP ──────────────────────────────────────────────────────────

def detectar_texto_eucjp(data: bytes) -> list:
    """Detecta texto EUC-JP (PS2/Saturn japoneses)."""
    bruto, n, candidatos, i = bytes(data), len(data), [], 0
    while i < n:
        inicio, chars = i, 0
        while i < n:
            b = bruto[i]
            if 0xA1 <= b <= 0xFE and i + 1 < n and 0xA1 <= bruto[i+1] <= 0xFE:
                try:
                    bruto[i:i+2].decode('euc_jp'); chars += 1; i += 2; continue
                except UnicodeDecodeError: pass
            if 0x20 <= b <= 0x7E: i += 1; continue
            break
        if chars >= 2:
            texto = bruto[inicio:i].decode('euc_jp', errors='ignore')
            candidatos.append((inicio, i, texto, 'eucjp'))
        if i == inicio: i += 1
    return candidatos

def truncar_eucjp(texto: str, max_bytes: int) -> bytes:
    melhor = b''
    for i in range(len(texto), -1, -1):
        enc = texto[:i].encode('euc_jp', errors='ignore')
        if len(enc) <= max_bytes: melhor = enc; break
    return melhor + b' ' * (max_bytes - len(melhor))

# ─── PARSER NARC (Nintendo Archive — NDS) ────────────────────────────────────

def parsear_narc(data: bytes) -> list | None:
    try:
        if data[:4] != b'NARC': return None
        btaf_off    = 16
        if data[btaf_off:btaf_off+4] != b'BTAF': return None
        btaf_size   = int.from_bytes(data[btaf_off+4:btaf_off+8], 'little')
        file_count  = int.from_bytes(data[btaf_off+8:btaf_off+10], 'little')
        entradas = []
        for i in range(file_count):
            base  = btaf_off + 12 + i * 8
            start = int.from_bytes(data[base:base+4], 'little')
            end   = int.from_bytes(data[base+4:base+8], 'little')
            entradas.append((start, end))
        gmif_off = btaf_off + btaf_size
        if data[gmif_off:gmif_off+4] == b'BTNF':
            gmif_off += int.from_bytes(data[gmif_off+4:gmif_off+8], 'little')
        if data[gmif_off:gmif_off+4] != b'GMIF': return None
        dados_off = gmif_off + 8
        return [data[dados_off+s:dados_off+e] for s, e in entradas]
    except Exception: return None

def remontar_narc(original: bytes, arquivos_novos: list) -> bytes:
    try:
        btaf_off   = 16
        btaf_size  = int.from_bytes(original[btaf_off+4:btaf_off+8], 'little')
        gmif_off   = btaf_off + btaf_size
        btnf_chunk = b''
        if original[gmif_off:gmif_off+4] == b'BTNF':
            bsize = int.from_bytes(original[gmif_off+4:gmif_off+8], 'little')
            btnf_chunk = original[gmif_off:gmif_off+bsize]; gmif_off += bsize
        novos_dados = b''.join(arquivos_novos)
        pad = (4 - len(novos_dados) % 4) % 4
        novos_dados += b'\xFF' * pad
        gmif = b'GMIF' + (len(novos_dados)+8).to_bytes(4,'little') + novos_dados
        nova_btaf = bytearray(original[btaf_off:btaf_off+12])
        off = 0
        for arq in arquivos_novos:
            nova_btaf += off.to_bytes(4,'little'); off += len(arq)
            nova_btaf += off.to_bytes(4,'little')
        nova_btaf[4:8] = len(nova_btaf).to_bytes(4,'little')
        corpo  = bytes(nova_btaf) + btnf_chunk + gmif
        header = bytearray(original[:16])
        header[8:12] = (16+len(corpo)).to_bytes(4,'little')
        return bytes(header) + corpo
    except Exception: return original

# ─── PARSER BMG (Binary MeSsaGe — NDS) ───────────────────────────────────────

def parsear_bmg(data: bytes) -> tuple | None:
    try:
        if data[:8] != b'MESGbmg1': return None
        enc_byte = data[0x0E]
        encoding = {1:'cp1252', 2:'utf-16-le', 3:'shift_jis', 4:'utf-8'}.get(enc_byte,'utf-8')
        n_sec = int.from_bytes(data[0x0C:0x10], 'little')
        pos, dat1_off, inf1_off = 0x20, None, None
        for _ in range(n_sec):
            if pos + 8 > len(data): break
            magic = data[pos:pos+4]
            size  = int.from_bytes(data[pos+4:pos+8], 'little')
            if magic == b'INF1': inf1_off = pos
            if magic == b'DAT1': dat1_off = pos
            pos += size; pos = (pos+3) & ~3
        if dat1_off is None or inf1_off is None: return None
        entry_count = int.from_bytes(data[inf1_off+8:inf1_off+10], 'little')
        entry_size  = int.from_bytes(data[inf1_off+10:inf1_off+12], 'little')
        dat1_base   = dat1_off + 8
        strings = []
        for i in range(entry_count):
            entry_off = inf1_off + 16 + i * entry_size
            str_off   = int.from_bytes(data[entry_off:entry_off+4], 'little')
            raw = data[dat1_base+str_off:]
            if encoding == 'utf-16-le':
                end = 0
                while end+1 < len(raw) and (raw[end]!=0 or raw[end+1]!=0): end += 2
                s = raw[:end].decode('utf-16-le', errors='ignore')
            else:
                end = raw.find(b'\x00')
                s = raw[:end].decode(encoding, errors='ignore') if end != -1 else ''
            strings.append(s)
        return strings, data
    except Exception: return None

def remontar_bmg(original: bytes, strings_novas: list) -> bytes:
    try:
        enc_byte = original[0x0E]
        encoding = {1:'cp1252', 2:'utf-16-le', 3:'shift_jis', 4:'utf-8'}.get(enc_byte,'utf-8')
        n_sec = int.from_bytes(original[0x0C:0x10], 'little')
        pos, dat1_off, inf1_off = 0x20, None, None
        for _ in range(n_sec):
            if pos + 8 > len(original): break
            magic = original[pos:pos+4]
            size  = int.from_bytes(original[pos+4:pos+8], 'little')
            if magic == b'INF1': inf1_off = pos
            if magic == b'DAT1': dat1_off = pos
            pos += size; pos = (pos+3) & ~3
        if dat1_off is None or inf1_off is None: return original
        entry_count = int.from_bytes(original[inf1_off+8:inf1_off+10], 'little')
        entry_size  = int.from_bytes(original[inf1_off+10:inf1_off+12], 'little')
        dat1_base   = dat1_off + 8
        dat1_size   = int.from_bytes(original[dat1_off+4:dat1_off+8], 'little') - 8
        novo_dat, novos_offsets = bytearray(), []
        for i, s in enumerate(strings_novas[:entry_count]):
            novos_offsets.append(len(novo_dat))
            enc_s = (s.encode('utf-16-le',errors='ignore')+b'\x00\x00') if encoding=='utf-16-le' else (s.encode(encoding,errors='ignore')+b'\x00')
            novo_dat += enc_s
        while len(novo_dat) % 4 != 0: novo_dat += b'\x00'
        result = bytearray(original)
        for i, off in enumerate(novos_offsets):
            addr = inf1_off + 16 + i * entry_size
            result[addr:addr+4] = off.to_bytes(4,'little')
        result[dat1_base:dat1_base+min(len(novo_dat),dat1_size)] = bytes(novo_dat)[:dat1_size]
        return bytes(result)
    except Exception: return original

# ─── PARSER MSBT (Message Studio Binary Text — 3DS/Switch) ───────────────────

def parsear_msbt(data: bytes) -> list | None:
    try:
        if data[:8] != b'MsgStdBn': return None
        little = data[8:10] == b'\xFF\xFE'
        bo, enc = ('little','utf-16-le') if little else ('big','utf-16-be')
        n_sec = int.from_bytes(data[0x0E:0x10], bo)
        pos, txt2_off = 0x20, None
        for _ in range(n_sec):
            if pos + 8 > len(data): break
            magic = data[pos:pos+4]; size = int.from_bytes(data[pos+4:pos+8], bo)
            if magic == b'TXT2': txt2_off = pos
            pos += size; pos = (pos+3) & ~3
        if txt2_off is None: return None
        count = int.from_bytes(data[txt2_off+8:txt2_off+12], bo)
        base  = txt2_off + 8
        strings = []
        for i in range(count):
            off_addr = base + 4 + i * 4
            str_off  = int.from_bytes(data[off_addr:off_addr+4], bo)
            raw = data[base + str_off:]
            end = 0
            while end+1 < len(raw) and (raw[end]!=0 or raw[end+1]!=0): end += 2
            s = raw[:end].decode(enc, errors='ignore')
            s = re.sub(r'\x0e.{3}', '', s).strip()
            if s: strings.append(s)
        return strings
    except Exception: return None

def remontar_msbt(original: bytes, strings_novas: list) -> bytes:
    try:
        little = original[8:10] == b'\xFF\xFE'
        bo, enc = ('little','utf-16-le') if little else ('big','utf-16-be')
        n_sec = int.from_bytes(original[0x0E:0x10], bo)
        pos, txt2_off = 0x20, None
        for _ in range(n_sec):
            if pos + 8 > len(original): break
            magic = original[pos:pos+4]; size = int.from_bytes(original[pos+4:pos+8], bo)
            if magic == b'TXT2': txt2_off = pos
            pos += size; pos = (pos+3) & ~3
        if txt2_off is None: return original
        txt2_size = int.from_bytes(original[txt2_off+4:txt2_off+8], bo) - 8
        count = int.from_bytes(original[txt2_off+8:txt2_off+12], bo)
        base  = txt2_off + 8
        header_size = 4 + count * 4
        novo_dat, novos_offs = bytearray(), []
        for s in strings_novas[:count]:
            novos_offs.append(header_size + len(novo_dat))
            novo_dat += s.encode(enc, errors='ignore') + b'\x00\x00'
        while len(novo_dat) % 4 != 0: novo_dat += b'\x00'
        result = bytearray(original)
        for i, off in enumerate(novos_offs):
            addr = base + 4 + i * 4
            result[addr:addr+4] = off.to_bytes(4, bo)
        espaco = txt2_size - header_size
        result[base+header_size:base+header_size+espaco] = bytes(novo_dat)[:espaco]
        return bytes(result)
    except Exception: return original

# ─── PARSER SARC (Simple ARChive — 3DS/Switch) ───────────────────────────────

def parsear_sarc(data: bytes) -> dict | None:
    try:
        if data[:4] != b'SARC': return None
        little = data[6:8] == b'\xFF\xFE'
        bo = 'little' if little else 'big'
        header_len  = int.from_bytes(data[4:6], bo)
        data_offset = int.from_bytes(data[0x0C:0x10], bo)
        sfat_off    = header_len
        if data[sfat_off:sfat_off+4] != b'SFAT': return None
        sfat_hlen  = int.from_bytes(data[sfat_off+4:sfat_off+6], bo)
        node_count = int.from_bytes(data[sfat_off+8:sfat_off+10], bo)
        nodes = []
        for i in range(node_count):
            base  = sfat_off + sfat_hlen + i * 16
            attrs = int.from_bytes(data[base+4:base+8], bo)
            start = int.from_bytes(data[base+8:base+12], bo)
            end   = int.from_bytes(data[base+12:base+16], bo)
            nodes.append((attrs, start, end))
        sfnt_off   = sfat_off + sfat_hlen + node_count * 16
        sfnt_hlen  = int.from_bytes(data[sfnt_off+4:sfnt_off+6], bo)
        names_base = sfnt_off + sfnt_hlen
        arquivos = {}
        for (attrs, start, end) in nodes:
            name_off = (attrs & 0xFFFF) * 4
            nome_raw = data[names_base+name_off:]
            nome = nome_raw[:nome_raw.find(b'\x00')].decode('utf-8', errors='ignore')
            arquivos[nome] = data[data_offset+start:data_offset+end]
        return arquivos
    except Exception: return None

def remontar_sarc(original: bytes, arquivos_novos: dict) -> bytes:
    try:
        little = original[6:8] == b'\xFF\xFE'
        bo = 'little' if little else 'big'
        header_len  = int.from_bytes(original[4:6], bo)
        data_offset = int.from_bytes(original[0x0C:0x10], bo)
        sfat_off    = header_len
        sfat_hlen   = int.from_bytes(original[sfat_off+4:sfat_off+6], bo)
        node_count  = int.from_bytes(original[sfat_off+8:sfat_off+10], bo)
        sfnt_off    = sfat_off + sfat_hlen + node_count * 16
        sfnt_hlen   = int.from_bytes(original[sfnt_off+4:sfnt_off+6], bo)
        names_base  = sfnt_off + sfnt_hlen
        result = bytearray(original)
        for i in range(node_count):
            base  = sfat_off + sfat_hlen + i * 16
            attrs = int.from_bytes(original[base+4:base+8], bo)
            start = int.from_bytes(original[base+8:base+12], bo)
            end   = int.from_bytes(original[base+12:base+16], bo)
            name_off = (attrs & 0xFFFF) * 4
            nome_raw = original[names_base+name_off:]
            nome = nome_raw[:nome_raw.find(b'\x00')].decode('utf-8', errors='ignore')
            if nome in arquivos_novos:
                novo   = arquivos_novos[nome]
                espaco = end - start
                result[data_offset+start:data_offset+start+espaco] = bytes(novo)[:espaco]
        return bytes(result)
    except Exception: return original


def detectar_texto_shiftjis(data):
    bruto = bytes(data)
    n = len(bruto)
    candidatos = []
    i = 0
    while i < n:
        inicio = i
        chars = []
        while i < n:
            b = bruto[i]
            if (0x81 <= b <= 0x9F or 0xE0 <= b <= 0xFC) and i + 1 < n:
                b2 = bruto[i + 1]
                if (0x40 <= b2 <= 0x7E) or (0x80 <= b2 <= 0xFC):
                    try:
                        bruto[i:i + 2].decode('shift_jis')
                        chars.append(True)
                        i += 2
                        continue
                    except UnicodeDecodeError:
                        pass
            if 0x20 <= b <= 0x7E or 0xA1 <= b <= 0xDF:
                chars.append(False)
                i += 1
                continue
            break
        if sum(chars) >= 2:
            texto = bruto[inicio:i].decode('shift_jis', errors='ignore')
            candidatos.append((inicio, i, texto, 'shiftjis'))
        if i == inicio:
            i += 1
    return candidatos

def remover_sobreposicoes(prioritarios, outros):
    faixas = [(c[0], c[1]) for c in prioritarios]
    def sobrepoe(inicio, fim):
        return any(not (fim <= s or inicio >= e) for s, e in faixas)
    filtrados = [c for c in outros if not sobrepoe(c[0], c[1])]
    return prioritarios + filtrados

def truncar_sjis(texto, max_bytes):
    limite_superior = min(len(texto), max_bytes)
    melhor = b''
    for i in range(limite_superior, -1, -1):
        codificado = texto[:i].encode('shift_jis', errors='ignore')
        if len(codificado) <= max_bytes:
            melhor = codificado
            break
    return melhor + b' ' * (max_bytes - len(melhor))

def processar_binario(caminho):
    with open(caminho, 'rb') as f:
        data = bytearray(f.read())

    # ── SEM tentativa de descompressão aqui ──────────────────────────────────
    # LZ só é descomprimido dentro de processar_nds via ndspy,
    # que sabe remontar o arquivo corretamente.
    # Aqui fazemos varredura direta no binário original.

    shiftjis_cands = detectar_texto_shiftjis(data)
    eucjp_cands    = detectar_texto_eucjp(data)
    outros_cands   = detectar_texto_8bit(data) + detectar_texto_16bit(data)

    cands_jp   = remover_sobreposicoes(shiftjis_cands, eucjp_cands)
    candidatos = remover_sobreposicoes(cands_jp, outros_cands)
    candidatos.sort(key=lambda c: c[0])

    comprimidos = detectar_blocos_comprimidos(data)
    if comprimidos:
        print(f"AVISO: {len(comprimidos)} bloco(s) comprimido(s) detectados. Exemplo: \"{comprimidos[0][1]}\"", flush=True)

    if not candidatos:
        raise Exception("Nenhum texto legível encontrado neste arquivo binário.")

    n_8     = sum(1 for c in candidatos if c[3] == '8bit')
    n_16    = sum(1 for c in candidatos if c[3] == '16bit')
    n_sjis  = sum(1 for c in candidatos if c[3] == 'shiftjis')
    n_eucjp = sum(1 for c in candidatos if c[3] == 'eucjp')
    print(f"Encontrados {len(candidatos)} trecho(s) ({n_8} 8bit, {n_16} 16bit, {n_sjis} ShiftJIS, {n_eucjp} EUC-JP). Traduzindo...", flush=True)

    indices_normais = [i for i, c in enumerate(candidatos) if c[3] in ('8bit','16bit')]
    indices_sjis    = [i for i, c in enumerate(candidatos) if c[3] == 'shiftjis']
    indices_eucjp   = [i for i, c in enumerate(candidatos) if c[3] == 'eucjp']
    traduzidos = [''] * len(candidatos)

    amostra_normal = ' '.join(candidatos[i][2] for i in indices_normais[:10])
    idioma_bin = detectar_idioma_texto(amostra_normal) if amostra_normal else 'en'
    dict_key_normal = get_dict_key(idioma_bin, _idioma_global)
    dict_key_jp     = get_dict_key('ja', _idioma_global)
    print(f"📖 Binário: idioma {idioma_bin} → {dict_key_normal}", flush=True)

    INSTRUCAO_BIN = (
        "Estes são textos de um arquivo binário de jogo (diálogos, menus, itens). "
        "REGRA CRÍTICA: cada tradução deve ter NO MÁXIMO o mesmo número de caracteres do original. "
        "Evite acentos. Responda apenas com o texto traduzido."
    )
    INSTRUCAO_JP = (
        "Textos japoneses de jogo (kanji/kana). "
        "REGRA CRÍTICA: seja extremamente conciso — kanji ocupa 2 bytes, português ocupa mais. "
        "Priorize a ideia central. Responda apenas com o texto traduzido."
    )

    if indices_normais:
        res = traduzir_lista([candidatos[i][2] for i in indices_normais], INSTRUCAO_BIN, dict_key=dict_key_normal)
        for i, t in zip(indices_normais, res): traduzidos[i] = t

    if indices_sjis:
        res = traduzir_lista([candidatos[i][2] for i in indices_sjis], INSTRUCAO_JP, dict_key=dict_key_jp)
        for i, t in zip(indices_sjis, res): traduzidos[i] = t

    if indices_eucjp:
        res = traduzir_lista([candidatos[i][2] for i in indices_eucjp], INSTRUCAO_JP, dict_key=dict_key_jp)
        for i, t in zip(indices_eucjp, res): traduzidos[i] = t

    for (start, end, original, modo), trad in zip(candidatos, traduzidos):
        max_len = end - start
        if modo == '8bit':
            tb = trad.encode('ascii', errors='ignore')[:max_len]
            tb += b' ' * (max_len - len(tb))
        elif modo == '16bit':
            num_chars = max_len // 2
            tb = bytearray()
            for ch in trad.encode('ascii', errors='ignore')[:num_chars]: tb += bytes([ch, 0])
            while len(tb) < max_len: tb += b' \x00'
            tb = bytes(tb)
        elif modo == 'shiftjis':
            tb = truncar_sjis(trad, max_len)
        elif modo == 'eucjp':
            tb = truncar_eucjp(trad, max_len)
        else:
            tb = trad.encode('ascii', errors='ignore')[:max_len]
            tb += b' ' * (max_len - len(tb))
        data[start:end] = tb

    ext_original = os.path.splitext(caminho)[1].lower() or '.bin'
    return bytes(data), ext_original


def processar_iso(caminho):
    import pycdlib
    import shutil

    print("Tentando abrir ISO...", flush=True)

    try:
        iso = pycdlib.PyCdlib()
        pasta_test = os.path.dirname(caminho)
        nome_base_test = os.path.splitext(os.path.basename(caminho))[0]
        copia_path = os.path.join(pasta_test, f"{nome_base_test}_iso_work.iso")

        print("Copiando ISO para trabalho (pode demorar)...", flush=True)
        shutil.copy2(caminho, copia_path)
        iso.open(copia_path, mode='r+b')

        arquivos_iso = []

        def caminhar(path):
            for child in iso.list_children(iso_path=path):
                nome = child.file_identifier()
                nome_str = nome.decode('utf-8', errors='ignore') if isinstance(nome, bytes) else nome
                if nome_str in ('.', '..'):
                    continue
                caminho_filho = path.rstrip('/') + '/' + nome_str
                if child.is_dir():
                    caminhar(caminho_filho)
                else:
                    arquivos_iso.append(caminho_filho)

        caminhar('/')
        print(f"Encontrados {len(arquivos_iso)} arquivo(s) dentro do ISO.", flush=True)

        traduzidos_count = 0
        ignorados_grandes = 0

        for iso_path in arquivos_iso:
            buffer = io.BytesIO()
            try:
                iso.get_file_from_iso_fp(buffer, iso_path=iso_path)
            except Exception:
                continue

            dados_originais = buffer.getvalue()
            if not dados_originais or len(dados_originais) > MAX_TAMANHO_ISO_ARQUIVO:
                if len(dados_originais) > MAX_TAMANHO_ISO_ARQUIVO:
                    ignorados_grandes += 1
                continue

            nome_limpo = iso_path.split(';')[0]
            temp_path = os.path.join(tempfile.gettempdir(), f"_campfire_iso_tmp{os.path.splitext(nome_limpo)[1] or '.bin'}")
            with open(temp_path, 'wb') as f:
                f.write(dados_originais)

            try:
                novo_conteudo, _ = processar_binario(temp_path)
            except Exception:
                os.remove(temp_path)
                continue
            os.remove(temp_path)

            if len(novo_conteudo) != len(dados_originais):
                continue

            try:
                iso.modify_file_in_place(io.BytesIO(novo_conteudo), len(novo_conteudo), iso_path)
                print(f"Reescrevendo: {nome_limpo}", flush=True)
                traduzidos_count += 1
            except Exception as e:
                print(f"AVISO: não foi possível reescrever {nome_limpo}: {e}", flush=True)

        iso.close()
        print(f"Arquivos traduzidos no ISO: {traduzidos_count}. Ignorados por tamanho: {ignorados_grandes}.", flush=True)

        with open(copia_path, 'rb') as f:
            resultado = f.read()
        os.remove(copia_path)
        return resultado, '.iso'

    except Exception as e:
        print(f"AVISO: pycdlib não conseguiu abrir este ISO ({e}).", flush=True)
        print("Usando modo de varredura direta por blocos. Pode demorar alguns minutos...", flush=True)

        work_path = caminho.replace('.iso', '_iso_work.iso')
        if os.path.exists(work_path):
            os.remove(work_path)

        CHUNK_SIZE = 5 * 1024 * 1024
        OVERLAP = 512
        tamanho_total = os.path.getsize(caminho)
        tamanho_mb = tamanho_total / 1_000_000
        totalBlocos = (tamanho_total + CHUNK_SIZE - 1) // CHUNK_SIZE
        print(f"Tamanho do ISO: {tamanho_mb:.0f}MB. Varrendo em blocos de 5MB...", flush=True)

        todos_candidatos = []
        offset_global = 0
        bloco_num = 0

        with open(caminho, 'rb') as f:
            while offset_global < tamanho_total:
                inicio_leitura = max(0, offset_global - OVERLAP)
                f.seek(inicio_leitura)
                chunk = f.read(CHUNK_SIZE + OVERLAP)
                if not chunk:
                    break

                data_chunk = bytearray(chunk)
                cands_sjis = detectar_texto_shiftjis(data_chunk)
                outros = detectar_texto_8bit(data_chunk) + detectar_texto_16bit(data_chunk)
                cands = remover_sobreposicoes(cands_sjis, outros)

                for c in cands:
                    todos_candidatos.append((c[0] + inicio_leitura, c[1] + inicio_leitura, c[2], c[3]))

                bloco_num += 1
                progresso = min(offset_global + CHUNK_SIZE, tamanho_total) / 1_000_000
                print(f"Bloco {bloco_num}/{totalBlocos} varrido ({progresso:.0f}MB/{tamanho_mb:.0f}MB)...", flush=True)
                offset_global += CHUNK_SIZE

        todos_candidatos.sort(key=lambda c: c[0])
        unicos = []
        ultimo_fim = -1
        for c in todos_candidatos:
            if c[0] >= ultimo_fim:
                unicos.append(c)
                ultimo_fim = c[1]
        todos_candidatos = unicos

        if not todos_candidatos:
            raise Exception("Nenhum texto legível encontrado no ISO.")

        n_8 = sum(1 for c in todos_candidatos if c[3] == '8bit')
        n_16 = sum(1 for c in todos_candidatos if c[3] == '16bit')
        n_sjis = sum(1 for c in todos_candidatos if c[3] == 'shiftjis')
        print(f"Total: {len(todos_candidatos)} trecho(s) ({n_8} em 8-bit, {n_16} em 16-bit, {n_sjis} em japonês). Traduzindo...", flush=True)

        indices_normais = [i for i, c in enumerate(todos_candidatos) if c[3] != 'shiftjis']
        indices_sjis = [i for i, c in enumerate(todos_candidatos) if c[3] == 'shiftjis']
        traduzidos = [''] * len(todos_candidatos)

        if indices_normais:
            textos_normais = [todos_candidatos[i][2] for i in indices_normais]
            resultado_normais = traduzir_lista(
                textos_normais,
                "Estes são textos extraídos de um arquivo ISO de jogo (diálogos, menus, itens). "
                "REGRA CRÍTICA: cada tradução deve ter NO MÁXIMO o mesmo número de caracteres do texto original. "
                "Evite acentos. Responda apenas com o texto traduzido, sem comentários."
            )
            for i, trad in zip(indices_normais, resultado_normais):
                traduzidos[i] = trad

        if indices_sjis:
            textos_sjis = [todos_candidatos[i][2] for i in indices_sjis]
            resultado_sjis = traduzir_lista(
                textos_sjis,
                "Estes textos estão em japonês (kanji/kana), extraídos de um ISO de jogo de PS2. "
                "REGRA CRÍTICA: seja extremamente conciso, kanji é muito mais compacto que português. "
                "Priorize a ideia central. Responda apenas com o texto traduzido."
            )
            for i, trad in zip(indices_sjis, resultado_sjis):
                traduzidos[i] = trad

        print("Aplicando traduções no arquivo...", flush=True)
        import shutil
        saida_raw = os.path.join(
            os.path.dirname(caminho),
            os.path.splitext(os.path.basename(caminho))[0] + '_traduzido_raw.iso'
        )
        shutil.copy2(caminho, saida_raw)

        with open(saida_raw, 'r+b') as f:
            for (start, end, original, modo), trad in zip(todos_candidatos, traduzidos):
                max_len = end - start
                if modo == '8bit':
                    trad_bytes = trad.encode('ascii', errors='ignore')[:max_len]
                    trad_bytes += b' ' * (max_len - len(trad_bytes))
                elif modo == '16bit':
                    num_chars = max_len // 2
                    trad_ascii = trad.encode('ascii', errors='ignore')[:num_chars]
                    trad_bytes = bytearray()
                    for ch in trad_ascii:
                        trad_bytes += bytes([ch, 0])
                    while len(trad_bytes) < max_len:
                        trad_bytes += b' \x00'
                    trad_bytes = bytes(trad_bytes)
                else:
                    trad_bytes = truncar_sjis(trad, max_len)
                f.seek(start)
                f.write(trad_bytes)

        print("Raw scan concluído.", flush=True)
        with open(saida_raw, 'rb') as f:
            resultado = f.read()
        os.remove(saida_raw)
        return resultado, '.iso'

def processar_zip(caminho):
    import zipfile

    pasta = os.path.dirname(caminho)
    nome_base = os.path.splitext(os.path.basename(caminho))[0]
    pasta_saida = os.path.join(pasta, f"{nome_base}_traduzido")
    os.makedirs(pasta_saida, exist_ok=True)

    with zipfile.ZipFile(caminho, 'r') as zf:
        arquivos = [n for n in zf.namelist() if not n.endswith('/')]

    print(f"Encontrados {len(arquivos)} arquivo(s) dentro do ZIP.", flush=True)

    traduzidos_count = 0
    ignorados = 0

    with zipfile.ZipFile(caminho, 'r') as zf:
        for nome_interno in arquivos:
            ext = os.path.splitext(nome_interno)[1].lower()

            if ext not in EXTENSOES or ext == '.zip':
                ignorados += 1
                continue

            dados = zf.read(nome_interno)
            temp_path = os.path.join(pasta_saida, f"_tmp_{os.path.basename(nome_interno)}")
            with open(temp_path, 'wb') as f:
                f.write(dados)

            try:
                conteudo, ext_saida = EXTENSOES[ext](temp_path)
                nome_saida = os.path.splitext(os.path.basename(nome_interno))[0] + f"_traduzido{ext_saida}"
                caminho_saida = os.path.join(pasta_saida, nome_saida)

                if isinstance(conteudo, bytes):
                    with open(caminho_saida, 'wb') as f: f.write(conteudo)
                else:
                    with open(caminho_saida, 'w', encoding='utf-8') as f: f.write(conteudo)

                print(f"OK: {nome_saida}", flush=True)
                traduzidos_count += 1
            except Exception as e:
                print(f"AVISO: {nome_interno} — {e}", flush=True)
            finally:
                if os.path.exists(temp_path):
                    os.remove(temp_path)

    print(f"ZIP concluído: {traduzidos_count} traduzido(s), {ignorados} ignorado(s).", flush=True)
    print(f"Arquivos salvos em: {pasta_saida}", flush=True)
    return b'', '.zip_done'

def processar_rar(caminho):
    import rarfile

    pasta = os.path.dirname(caminho)
    nome_base = os.path.splitext(os.path.basename(caminho))[0]
    pasta_saida = os.path.join(pasta, f"{nome_base}_traduzido")
    os.makedirs(pasta_saida, exist_ok=True)

    with rarfile.RarFile(caminho, 'r') as rf:
        arquivos = [n for n in rf.namelist() if not n.endswith('/')]

    print(f"Encontrados {len(arquivos)} arquivo(s) dentro do RAR.", flush=True)

    traduzidos_count = 0
    ignorados = 0

    with rarfile.RarFile(caminho, 'r') as rf:
        for nome_interno in arquivos:
            ext = os.path.splitext(nome_interno)[1].lower()

            if ext not in EXTENSOES or ext in ('.zip', '.rar'):
                ignorados += 1
                continue

            dados = rf.read(nome_interno)
            temp_path = os.path.join(pasta_saida, f"_tmp_{os.path.basename(nome_interno)}")
            with open(temp_path, 'wb') as f:
                f.write(dados)

            try:
                conteudo, ext_saida = EXTENSOES[ext](temp_path)
                nome_saida = os.path.splitext(os.path.basename(nome_interno))[0] + f"_traduzido{ext_saida}"
                caminho_saida = os.path.join(pasta_saida, nome_saida)

                if isinstance(conteudo, bytes):
                    with open(caminho_saida, 'wb') as f: f.write(conteudo)
                else:
                    with open(caminho_saida, 'w', encoding='utf-8') as f: f.write(conteudo)

                print(f"OK: {nome_saida}", flush=True)
                traduzidos_count += 1
            except Exception as e:
                print(f"AVISO: {nome_interno} — {e}", flush=True)
            finally:
                if os.path.exists(temp_path):
                    os.remove(temp_path)

    print(f"RAR concluído: {traduzidos_count} traduzido(s), {ignorados} ignorado(s).", flush=True)
    print(f"Arquivos salvos em: {pasta_saida}", flush=True)
    return b'', '.rar_done'

# ─── IMAGENS / PAINÉIS ───────────────────────────────────────────────────────

# ─── IMAGENS / PAINÉIS ───────────────────────────────────────────────────────

def processar_imagem(caminho):
    """
    Usa Claude Vision pra detectar texto com coordenadas, desenha caixas
    numeradas coloridas na imagem e gera um painel de legendas embaixo.
    Contribui os pares original→tradução pro dicionário colaborativo.
    """
    import base64
    import json as _json
    from PIL import Image, ImageDraw, ImageFont
    import io as _io

    TIPOS_MIME = {
        '.jpg':  'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.png':  'image/png',
        '.webp': 'image/webp',
        '.gif':  'image/gif',
    }

    ext = os.path.splitext(caminho)[1].lower()
    mime = TIPOS_MIME.get(ext, 'image/jpeg')

    with open(caminho, 'rb') as f:
        img_bytes = f.read()
    img_b64 = base64.b64encode(img_bytes).decode('utf-8')

    img_original = Image.open(_io.BytesIO(img_bytes)).convert('RGBA')
    w, h = img_original.size

    print("Analisando imagem com Claude Vision...", flush=True)

    msg = client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=4096,
        messages=[{
            "role": "user",
            "content": [
                {
                    "type": "image",
                    "source": {"type": "base64", "media_type": mime, "data": img_b64},
                },
                {
                    "type": "text",
                    "text": (
                        f"Analise esta imagem e encontre TODO o texto visível.\n\n"
                        f"Retorne SOMENTE um JSON válido, sem markdown, sem explicações:\n"
                        f"[{{\"n\":1,\"original\":\"texto aqui\",\"traducao\":\"tradução para {_idioma_global}\","
                        f"\"x1\":10,\"y1\":5,\"x2\":40,\"y2\":15}}]\n\n"
                        f"x1,y1 = canto superior esquerdo do texto em % da imagem (0-100).\n"
                        f"x2,y2 = canto inferior direito em % da imagem (0-100).\n"
                        f"Se não houver texto, retorne: []"
                    )
                }
            ],
        }]
    )

    resposta = msg.content[0].text.strip()

    try:
        itens = _json.loads(resposta)
    except Exception:
        match = re.search(r'\[.*\]', resposta, re.DOTALL)
        itens = _json.loads(match.group()) if match else []

    if not itens:
        raise Exception("Nenhum texto encontrado nesta imagem.")

    print(f"Encontrados {len(itens)} trecho(s) de texto. Montando imagem anotada...", flush=True)

    # ── Contribui pares pro dicionário ────────────────────────────────────────
    try:
        amostra = ' '.join(item['original'] for item in itens[:5])
        idioma_origem = detectar_idioma_texto(amostra)
        dict_key = get_dict_key(idioma_origem, _idioma_global)
        pares = {item['original']: item['traducao'] for item in itens
                 if item.get('original') and item.get('traducao')
                 and item['original'] != item['traducao']}
        if pares:
            contribute_dictionary(dict_key, pares)
            print(f"📖 Imagem: {len(pares)} par(es) contribuídos → dicionário: {dict_key}", flush=True)
    except Exception as e:
        print(f"⚠️ Erro ao contribuir imagem pro dicionário: {e}", flush=True)

    CORES = [
        (255, 59,  59),
        (59,  130, 255),
        (59,  200, 59),
        (255, 165, 0),
        (180, 59,  200),
        (59,  210, 210),
        (255, 220, 50),
    ]

    overlay = Image.new('RGBA', img_original.size, (0, 0, 0, 0))
    draw_ov = ImageDraw.Draw(overlay)

    for item in itens:
        cor = CORES[(item['n'] - 1) % len(CORES)]
        x1 = int(item['x1'] / 100 * w)
        y1 = int(item['y1'] / 100 * h)
        x2 = int(item['x2'] / 100 * w)
        y2 = int(item['y2'] / 100 * h)

        draw_ov.rectangle([x1, y1, x2, y2], fill=(*cor, 50), outline=(*cor, 230), width=2)
        draw_ov.rectangle([x1, y1, x1 + 18, y1 + 18], fill=(*cor, 230))
        draw_ov.text((x1 + 3, y1 + 1), str(item['n']), fill=(255, 255, 255, 255))

    img_anotada = Image.alpha_composite(img_original, overlay).convert('RGB')

    LINHA_H = 20
    PADDING = 14
    MAX_CHARS_POR_LINHA = 80

    try:
        fonte_titulo = ImageFont.truetype(r"C:\Windows\Fonts\arialbd.ttf", 14)
        fonte_texto  = ImageFont.truetype(r"C:\Windows\Fonts\arial.ttf",   13)
    except Exception:
        fonte_titulo = ImageFont.load_default()
        fonte_texto  = fonte_titulo

    def quebrar_linha(texto, max_chars):
        linhas = []
        while len(texto) > max_chars:
            corte = texto[:max_chars].rfind(' ')
            if corte == -1: corte = max_chars
            linhas.append(texto[:corte])
            texto = texto[corte:].strip()
        linhas.append(texto)
        return linhas

    linhas_por_item = []
    for item in itens:
        original_curto = item['original'] if len(item['original']) <= 40 else item['original'][:37] + '...'
        linha_completa = f"{original_curto}  →  {item['traducao']}"
        linhas_por_item.append(quebrar_linha(linha_completa, MAX_CHARS_POR_LINHA))

    total_linhas = sum(len(l) for l in linhas_por_item)
    painel_h = PADDING * 2 + 28 + total_linhas * LINHA_H + len(itens) * 6
    painel   = Image.new('RGB', (w, painel_h), (28, 28, 28))
    draw_p   = ImageDraw.Draw(painel)

    draw_p.text((PADDING, PADDING), "TRADUÇÕES", fill=(255, 153, 0), font=fonte_titulo)

    y_atual = PADDING + 26
    for item, linhas in zip(itens, linhas_por_item):
        cor = CORES[(item['n'] - 1) % len(CORES)]
        draw_p.rectangle([PADDING, y_atual + 2, PADDING + 18, y_atual + 18], fill=cor)
        draw_p.text((PADDING + 3, y_atual + 3), str(item['n']), fill=(255, 255, 255), font=fonte_texto)
        draw_p.text((PADDING + 22, y_atual + 3), linhas[0], fill=(220, 220, 220), font=fonte_texto)
        y_atual += LINHA_H
        for linha_extra in linhas[1:]:
            draw_p.text((PADDING + 22, y_atual + 3), linha_extra, fill=(180, 180, 180), font=fonte_texto)
            y_atual += LINHA_H
        y_atual += 6

    final = Image.new('RGB', (w, h + painel_h), (28, 28, 28))
    final.paste(img_anotada, (0, 0))
    final.paste(painel, (0, h))

    buf = _io.BytesIO()
    final.save(buf, format='PNG')
    return buf.getvalue(), '.png'
# ─── FORMATOS NINTENDO ────────────────────────────────────────────────────────

def processar_nds(caminho):
    """NDS com suporte a NARC, BMG, LZ10/LZ11 e varredura binária."""
    try:
        import ndspy.rom
        print("Abrindo ROM Nintendo DS com ndspy...", flush=True)
        rom = ndspy.rom.NintendoDSRom.fromFile(caminho)

        nome_por_indice = {}
        def _caminhar(fs, prefix=''):
            for nome, conteudo in fs.items():
                if isinstance(conteudo, dict): _caminhar(conteudo, prefix+nome+'/')
                else: nome_por_indice[conteudo] = prefix+nome
        if rom.filenames: _caminhar(rom.filenames)

        print(f"ROM: {len(rom.files)} arquivo(s) no filesystem.", flush=True)
        EXTS_NDS = {'', '.bin', '.dat', '.msg', '.txt', '.arc', '.narc', '.bmg', '.mtx'}
        dict_key = get_dict_key('ja', _idioma_global)
        traduzidos_count = 0

        for idx, dados in enumerate(rom.files):
            if not dados or len(dados) < 8: continue
            nome = nome_por_indice.get(idx, f'file_{idx:04d}.bin')
            ext  = os.path.splitext(nome)[1].lower()
            if ext not in EXTS_NDS: continue

            # Tenta descomprimir LZ
            descomp = tentar_descomprimir_nintendo(bytes(dados))
            dw = descomp if descomp and len(descomp) > 8 else bytes(dados)

            # Tenta NARC
            narc_files = parsear_narc(dw)
            if narc_files:
                novos, mod = list(narc_files), False
                for j, arq in enumerate(narc_files):
                    if not arq or len(arq) < 4: continue
                    bmg = parsear_bmg(arq)
                    if bmg:
                        strings, _ = bmg
                        sf = [s for s in strings if s.strip() and len(s) > 1]
                        if sf:
                            trad = traduzir_lista(sf, "Texto Nintendo DS. Conciso.", dict_key=dict_key)
                            mapa = {o:t for o,t in zip(sf,trad)}
                            novos[j] = remontar_bmg(arq, [mapa.get(s,s) for s in strings])
                            mod = True
                    else:
                        tmp = tempfile.mktemp(suffix='.bin')
                        with open(tmp,'wb') as f: f.write(arq)
                        try:
                            novo, _ = processar_binario(tmp)
                            if novo != arq: novos[j] = novo; mod = True
                        except Exception: pass
                        finally:
                            if os.path.exists(tmp): os.remove(tmp)
                if mod:
                    rom.files[idx] = remontar_narc(dw, novos)
                    traduzidos_count += 1
                    print(f"NARC traduzido: {nome}", flush=True)
                continue

            # Tenta BMG direto
            bmg = parsear_bmg(dw)
            if bmg:
                strings, _ = bmg
                sf = [s for s in strings if s.strip() and len(s) > 1]
                if sf:
                    trad = traduzir_lista(sf, "Texto Nintendo DS. Conciso.", dict_key=dict_key)
                    mapa = {o:t for o,t in zip(sf,trad)}
                    rom.files[idx] = remontar_bmg(dw, [mapa.get(s,s) for s in strings])
                    traduzidos_count += 1
                    print(f"BMG traduzido: {nome}", flush=True)
                continue

            # Fallback binário
            tmp = tempfile.mktemp(suffix=ext or '.bin')
            with open(tmp,'wb') as f: f.write(dw)
            try:
                novo, _ = processar_binario(tmp)
                if novo != dw:
                    rom.files[idx] = novo; traduzidos_count += 1
                    print(f"Binário traduzido: {nome}", flush=True)
            except Exception: pass
            finally:
                if os.path.exists(tmp): os.remove(tmp)

        print(f"Total NDS: {traduzidos_count} arquivo(s) traduzido(s).", flush=True)
        return bytes(rom.save()), '.nds'

    except ImportError:
        print("ndspy não instalado. Varredura binária...", flush=True)
        return processar_binario(caminho)
    except Exception as e:
        print(f"Erro ndspy ({e}). Varredura binária...", flush=True)
        return processar_binario(caminho)


def _processar_nintendo_sarc_scan(caminho: str, ext_saida: str) -> tuple:
    """Varredura de SARC/MSBT dentro de ROMs 3DS/Switch."""
    with open(caminho, 'rb') as f:
        data = f.read()

    dict_key = get_dict_key('ja', _idioma_global)
    result   = bytearray(data)
    count    = 0
    pos      = 0

    while pos < len(data) - 4:
        if data[pos:pos+4] == b'SARC':
            try:
                little    = data[pos+6:pos+8] == b'\xFF\xFE'
                bo        = 'little' if little else 'big'
                sarc_size = int.from_bytes(data[pos+8:pos+12], bo)
                if 0 < sarc_size <= len(data) - pos:
                    sarc_data = data[pos:pos+sarc_size]
                    arquivos  = parsear_sarc(sarc_data)
                    if arquivos:
                        mod, arqs_novos = False, dict(arquivos)
                        for nome, conteudo in arquivos.items():
                            strings = parsear_msbt(conteudo)
                            if strings:
                                sf = [s for s in strings if s.strip() and len(s) > 1]
                                if sf:
                                    trad = traduzir_lista(sf, "Texto Nintendo. Conciso.", dict_key=dict_key)
                                    arqs_novos[nome] = remontar_msbt(conteudo, trad)
                                    mod = True; count += 1
                                    print(f"MSBT traduzido: {nome}", flush=True)
                        if mod:
                            novo_sarc = remontar_sarc(sarc_data, arqs_novos)
                            result[pos:pos+len(novo_sarc)] = novo_sarc
                        pos += sarc_size; continue
            except Exception: pass
        pos += 1
        if pos % 50_000_000 == 0:
            print(f"Varrendo: {pos//1_000_000}MB...", flush=True)

    if count > 0:
        print(f"Total: {count} MSBT(s) traduzido(s).", flush=True)
        return bytes(result), ext_saida

    print("Nenhum SARC/MSBT encontrado. Varredura binária...", flush=True)
    return processar_binario(caminho)


def processar_3ds(caminho):
    """3DS: busca SARC/MSBT (ROMs descriptografadas) + fallback binário."""
    print("Processando Nintendo 3DS...", flush=True)
    print("ℹ️ Para melhor resultado, use ROM descriptografada (GodMode9).", flush=True)
    return _processar_nintendo_sarc_scan(caminho, '.3ds')


def processar_nsp(caminho):
    """NSP: busca SARC/MSBT (descriptografado) + fallback binário."""
    print("Processando NSP (Nintendo Switch Package)...", flush=True)
    print("ℹ️ Para melhor resultado, use NSP descriptografado (prod.keys).", flush=True)
    return _processar_nintendo_sarc_scan(caminho, '.nsp')


def processar_xci(caminho):
    """XCI: busca SARC/MSBT (descriptografado) + fallback binário."""
    print("Processando XCI (Nintendo Switch Game Card)...", flush=True)
    print("ℹ️ Para melhor resultado, use XCI descriptografado (prod.keys).", flush=True)
    return _processar_nintendo_sarc_scan(caminho, '.xci')
# ─── MAPA DE EXTENSÕES ────────────────────────────────────────────────────────

EXTENSOES = {
    '.txt':  processar_txt,
    '.srt':  processar_srt,
    '.json': processar_json,
    '.xml':  processar_xml,
    '.csv':  processar_csv,
    '.pdf':  processar_pdf,
    '.mkv':  processar_video,
    '.mp4':  processar_video,
    '.mp3':  processar_mp3,
    '.epub': processar_epub,
    '.docx': processar_docx,
    '.xlsx': processar_xlsx,
    # ── Localização de software (i18n) ──
    '.po':     processar_po,
    '.strings': processar_strings,
    '.resx':   processar_resx,
    '.bin':  processar_binario,
    '.dat':  processar_binario,
    '.iso':  processar_iso,
    '.zip':  processar_zip,
    '.rar':  processar_rar,
    # ── Nintendo ──
    '.nds':  processar_nds,
    '.3ds':  processar_3ds,
    '.nsp':  processar_nsp,
    '.xci':  processar_xci,
    # ── Imagens ──
    '.jpg':  processar_imagem,
    '.jpeg': processar_imagem,
    '.png':  processar_imagem,
    '.webp': processar_imagem,
    '.gif':  processar_imagem,
}

# ─── ARQUIVOS AUXILIARES ──────────────────────────────────────────────────────

TUTORIAL_VLC_PT = """COMO ASSISTIR COM A LEGENDA TRADUZIDA
========================================

Você recebeu um arquivo de legenda traduzido (.srt). Veja como usá-lo:

MÉTODO FÁCIL (recomendado)

1. Encontre o arquivo de vídeo original na sua pasta.
2. Renomeie o arquivo de legenda traduzido para ter EXATAMENTE o mesmo nome do vídeo, mas terminando em ".srt".
   Exemplo: se o vídeo se chama "filme.mkv", a legenda deve se chamar "filme.srt".
3. Coloque os dois arquivos (vídeo e legenda) na mesma pasta.
4. Abra o vídeo normalmente. A legenda traduzida vai aparecer automaticamente!

MÉTODO MANUAL (se o método fácil não funcionar)

Esse método usa o VLC, um programa gratuito e muito popular para assistir vídeos.
Se você não tem o VLC, baixe em: https://www.videolan.org/vlc/

1. Abra o vídeo no VLC.
2. No menu superior, clique em "Legenda" (ou "Subtitle" em inglês).
3. Clique em "Adicionar arquivo de legenda..." (ou "Add Subtitle File...").
4. Selecione o arquivo de legenda traduzido (.srt).
5. Pronto! A legenda traduzida vai aparecer no vídeo.

---
Tutorial gerado automaticamente por Campfire 🔥
"""

AVISO_BACKUP_PT = """AVISO IMPORTANTE
========================================

Antes de substituir o arquivo original do jogo ou aplicativo pelo arquivo traduzido,
faça uma cópia de segurança (backup) do arquivo original em outro local.

Isso evita perda de progresso, saves corrompidos ou problemas ao reverter
a tradução caso algo não funcione como esperado.

Se o arquivo traduzido for um arquivo binário de jogo (.bin, .dat, .gxt ou similar),
teste sempre numa cópia primeiro e confirme que o jogo ainda abre normalmente
antes de substituir o arquivo definitivo dentro da instalação do jogo.

---
Aviso gerado automaticamente por Campfire 🔥
"""

GUIA_ISO_PT = """COMO USAR O ARQUIVO ISO TRADUZIDO
========================================

Você recebeu uma cópia traduzida do disco original (.iso). Veja como usá-la:

NO COMPUTADOR (Windows 8 ou mais recente)

1. Clique com o botão direito no arquivo .iso traduzido.
2. Escolha "Montar" (ou "Mount" em inglês).
3. O Windows vai abrir o conteúdo como se fosse um CD/DVD inserido.

EM EMULADORES DE CONSOLE

A maioria dos emuladores (como o PCSX2, para PS2) permite carregar o arquivo .iso
diretamente pelo menu de abrir jogo/disco, sem precisar gravar em mídia física.

GRAVANDO EM DISCO FÍSICO

Se preferir gravar em um CD/DVD de verdade, use um programa como o ImgBurn (gratuito)
e selecione a opção de gravar a partir de uma imagem ISO.

IMPORTANTE

Nem todo arquivo dentro do disco pôde ser traduzido — arquivos muito grandes (como vídeos,
músicas e modelos 3D) são deixados intactos de propósito, para não correr o risco de
corromper o disco. Apenas arquivos pequenos de texto/dados foram traduzidos.

---
Guia gerado automaticamente por Campfire 🔥
"""

def gerar_tutorial(idioma):
    if idioma == "português brasileiro coloquial":
        return TUTORIAL_VLC_PT
    return traduzir_bloco(TUTORIAL_VLC_PT, idioma,
        'Este é um tutorial de instalação. Mantenha nomes próprios (VLC), URLs, nomes de arquivo e exemplos exatamente como estão. Traduza apenas o texto explicativo.')

def gerar_aviso_backup(idioma):
    if idioma == "português brasileiro coloquial":
        return AVISO_BACKUP_PT
    return traduzir_bloco(AVISO_BACKUP_PT, idioma,
        'Este é um aviso de segurança curto. Traduza o texto explicativo mantendo a formatação e o tom de aviso.')

def gerar_guia_iso(idioma):
    if idioma == "português brasileiro coloquial":
        return GUIA_ISO_PT
    return traduzir_bloco(GUIA_ISO_PT, idioma,
        'Este é um guia de uso de arquivo. Mantenha nomes próprios (Windows, PCSX2, ImgBurn) exatamente como estão.')

# ─── LOTE (PASTA) ─────────────────────────────────────────────────────────────

NOMES_AUXILIARES = {"tutorial_instalacao.txt", "aviso_backup.txt", "como_usar_iso.txt"}

def processar_pasta(pasta):
    arquivos = []
    for nome in sorted(os.listdir(pasta)):
        caminho_completo = os.path.join(pasta, nome)
        if not os.path.isfile(caminho_completo): continue
        if nome in NOMES_AUXILIARES: continue
        nome_base, ext = os.path.splitext(nome)
        ext = ext.lower()
        if ext not in EXTENSOES: continue
        if nome_base.endswith('_traduzido'): continue
        arquivos.append((caminho_completo, ext))

    if not arquivos:
        print("ERRO: nenhum arquivo suportado encontrado na pasta.", flush=True)
        return

    print(f"Encontrados {len(arquivos)} arquivo(s) suportado(s).", flush=True)
    sucesso, falhas = 0, 0
    precisa_tutorial, precisa_aviso, precisa_guia_iso = False, False, False

    for idx, (caminho_arq, ext) in enumerate(arquivos, 1):
        nome = os.path.basename(caminho_arq)
        print(f"\n[{idx}/{len(arquivos)}] {nome}", flush=True)
        try:
            handler = EXTENSOES[ext]
            conteudo, ext_saida = handler(caminho_arq)
            nome_base = os.path.splitext(nome)[0]
            saida = os.path.join(pasta, f"{nome_base}_traduzido{ext_saida}")
            if isinstance(conteudo, bytes):
                with open(saida, 'wb') as f: f.write(conteudo)
            else:
                with open(saida, 'w', encoding='utf-8') as f: f.write(conteudo)
            if ext_saida == '.srt': precisa_tutorial = True
            elif ext_saida == '.iso': precisa_guia_iso = True
            elif ext_saida in ('.json', '.xml') or handler is processar_binario: precisa_aviso = True
            print(f"OK: {saida}", flush=True)
            sucesso += 1
        except Exception as e:
            print(f"ERRO em {nome}: {e}", flush=True)
            falhas += 1

    if precisa_tutorial:
        with open(os.path.join(pasta, "tutorial_instalacao.txt"), 'w', encoding='utf-8') as f:
            f.write(gerar_tutorial(_idioma_global))
        print("Tutorial gerado.", flush=True)
    if precisa_aviso:
        with open(os.path.join(pasta, "aviso_backup.txt"), 'w', encoding='utf-8') as f:
            f.write(gerar_aviso_backup(_idioma_global))
        print("Aviso gerado.", flush=True)
    if precisa_guia_iso:
        with open(os.path.join(pasta, "como_usar_iso.txt"), 'w', encoding='utf-8') as f:
            f.write(gerar_guia_iso(_idioma_global))
        print("Guia de ISO gerado.", flush=True)

    print(f"\nResumo: {sucesso} traduzido(s), {falhas} com erro. Cache: {stats_cache()} entradas.", flush=True)

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("ERRO: nenhum caminho informado.", flush=True)
        sys.exit(1)

    caminho = sys.argv[1]
    _idioma_global = sys.argv[2] if len(sys.argv) > 2 else "português brasileiro coloquial"
    api_key = sys.argv[3] if len(sys.argv) > 3 else None

    inicializar_cliente(api_key)
    print(f"Idioma alvo: {_idioma_global}", flush=True)
    print(f"Cache: {stats_cache()} entradas disponíveis.", flush=True)

    if os.path.isdir(caminho):
        print(f"Modo lote: {caminho}", flush=True)
        processar_pasta(caminho)
        print(f"CONCLUIDO:{caminho}", flush=True)
        sys.exit(0)

    ext = os.path.splitext(caminho)[1].lower()
    usou_binario = False

    if ext in EXTENSOES:
        handler = EXTENSOES[ext]
        if handler is processar_binario: usou_binario = True
    else:
        print(f"Extensão '{ext}' não reconhecida, tentando como arquivo binário...", flush=True)
        handler = processar_binario
        usou_binario = True

    print(f"Iniciando tradução de: {os.path.basename(caminho)}", flush=True)
    conteudo, ext_saida = handler(caminho)

    pasta = os.path.dirname(caminho)
    nome_base = os.path.splitext(os.path.basename(caminho))[0]
    saida = os.path.join(pasta, f"{nome_base}_traduzido{ext_saida}")

    if isinstance(conteudo, bytes):
        with open(saida, 'wb') as f: f.write(conteudo)
    else:
        with open(saida, 'w', encoding='utf-8') as f: f.write(conteudo)

    if ext_saida == '.srt':
        with open(os.path.join(pasta, "tutorial_instalacao.txt"), 'w', encoding='utf-8') as f:
            f.write(gerar_tutorial(_idioma_global))
        print("Tutorial gerado.", flush=True)
    elif ext_saida == '.iso':
        with open(os.path.join(pasta, "como_usar_iso.txt"), 'w', encoding='utf-8') as f:
            f.write(gerar_guia_iso(_idioma_global))
        print("Guia gerado.", flush=True)
    elif ext_saida in ('.json', '.xml') or usou_binario:
        with open(os.path.join(pasta, "aviso_backup.txt"), 'w', encoding='utf-8') as f:
            f.write(gerar_aviso_backup(_idioma_global))
        print("Aviso gerado.", flush=True)

    print(f"Cache: {stats_cache()} entradas salvas.", flush=True)
    print(f"CONCLUIDO:{saida}", flush=True)